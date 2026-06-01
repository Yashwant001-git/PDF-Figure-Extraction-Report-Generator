from docx import Document
from docx.shared import Inches
import os


class ReportGenerator:

    def __init__(self, output_dir="data/reports"):

        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_report(
        self,
        figures,
        pdf_name,
        report_name="figure_report.docx"
    ):

        doc = Document()

        # Cover Title
        doc.add_heading(
            "PDF Figure Extraction Report",
            level=0
        )

        doc.add_paragraph(
            f"Source PDF: {pdf_name}"
        )

        doc.add_paragraph(
            f"Total Figures Extracted: {len(figures)}"
        )

        doc.add_page_break()

        # Summary Section
        doc.add_heading(
            "Summary",
            level=1
        )

        table = doc.add_table(
            rows=1,
            cols=3
        )

        table.style = "Table Grid"

        header = table.rows[0].cells

        header[0].text = "Figure"
        header[1].text = "Page"
        header[2].text = "Caption"

        for idx, figure in enumerate(figures, start=1):

            row = table.add_row().cells

            row[0].text = f"Figure {idx}"
            row[1].text = str(figure["page"])

            caption = figure["caption"]

            if len(caption) > 60:
                caption = caption[:60] + "..."

            row[2].text = caption

        doc.add_page_break()

        # Detailed Figures
        doc.add_heading(
            "Detailed Figures",
            level=1
        )

        for idx, figure in enumerate(figures, start=1):

            doc.add_heading(
                f"Figure {idx}",
                level=2
            )

            doc.add_paragraph(
                f"Page Number: {figure['page']}"
            )

            doc.add_paragraph(
                f"Caption: {figure['caption']}"
            )

            image_path = figure["image_path"]

            if os.path.exists(image_path):

                try:
                    doc.add_picture(
                        image_path,
                        width=Inches(4.5)
                    )

                except Exception as e:

                    doc.add_paragraph(
                        f"Error loading image: {e}"
                    )

            doc.add_page_break()

        report_path = os.path.join(
            self.output_dir,
            report_name
        )

        doc.save(report_path)

        return report_path